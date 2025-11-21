
import sys
import os
from docx import Document

def extract_text_from_docx(docx_path, output_txt_path):
    print(f"Extracting text from {docx_path}...")
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract from tables as curriculum is often in tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                full_text.append(" | ".join(row_text))
        
        text_content = '\n'.join(full_text)
        
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
            
        print(f"Successfully extracted text to {output_txt_path}")
        return True
    except Exception as e:
        print(f"Error extracting text: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_docx.py <docx_path> <output_txt_path>")
        sys.exit(1)
        
    docx_path = sys.argv[1]
    output_path = sys.argv[2]
    
    extract_text_from_docx(docx_path, output_path)
