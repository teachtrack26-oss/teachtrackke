#!/usr/bin/env python3
"""
Extract text from Word documents (.docx files)
"""

import sys
from docx import Document

def extract_text_from_docx(file_path):
    """Extract text from a Word document"""
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return None

def main():
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <path_to_docx_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    print(f"Extracting text from: {file_path}")
    
    text = extract_text_from_docx(file_path)
    
    if text:
        # Create output filename
        output_file = file_path.replace('.docx', '_extracted.txt')
        
        # Save extracted text
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        print(f"Text extracted and saved to: {output_file}")
        
        # Also print first 2000 characters to console
        print("\n" + "="*80)
        print("PREVIEW (first 2000 characters):")
        print("="*80)
        print(text[:2000])
        if len(text) > 2000:
            print("\n... (content truncated, check the output file for full text)")
    else:
        print("Failed to extract text from the document")

if __name__ == "__main__":
    main()